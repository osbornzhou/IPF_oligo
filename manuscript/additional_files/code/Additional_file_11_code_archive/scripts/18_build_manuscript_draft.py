#!/usr/bin/env python
"""
Build an editable BMC Genomics-style manuscript draft from project outputs.

The document is intentionally a scientific draft, not a finished submission:
author details, institutional approvals, final reference metadata, and figure
layout still need human confirmation before submission.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parents[1]
OUT_DIR = PROJECT_DIR / "manuscript"
OUT_DIR.mkdir(parents=True, exist_ok=True)

DOCX_PATH = OUT_DIR / "ipf_oligo_ml_bmc_genomics_manuscript_draft.docx"
MD_PATH = OUT_DIR / "ipf_oligo_ml_bmc_genomics_manuscript_draft.md"


def read_csv(rel_path: str) -> pd.DataFrame:
    return pd.read_csv(PROJECT_DIR / rel_path)


def fmt_float(value: object, digits: int = 3) -> str:
    try:
        return f"{float(value):.{digits}f}"
    except Exception:
        return str(value)


def fmt_p(value: object) -> str:
    try:
        value = float(value)
    except Exception:
        return str(value)
    if value == 0:
        return "0"
    if value < 0.001:
        return f"{value:.2e}"
    return f"{value:.3f}"


expr_qc = read_csv("metadata/expression_matrix_qc.csv")
de_qc = read_csv("results/differential_expression/differential_expression_qc.csv")
robust_qc = read_csv("results/robust_candidates/robust_candidate_qc.csv")
robust_mrna = read_csv("results/robust_candidates/robust_mrna_candidates_strict.csv")
robust_mirna = read_csv("results/robust_candidates/robust_mirna_candidates_strict.csv")
axis_qc = read_csv("results/mirna_mrna_axes/mirna_mrna_axis_qc.csv").iloc[0]
axes = read_csv("results/mirna_mrna_axes/top100_robust_mirna_mrna_axes.csv")
enrich_qc = read_csv("results/enrichment/enrichment_triple_qc.csv")
enrich_strict = read_csv("results/enrichment/robust_mrna_strict_enrichment_significant_fdr0.05.csv")
ppi_qc = read_csv("results/ppi_network/string_ppi_triple_qc.csv")
ppi_hubs = read_csv("results/ppi_network/string_ppi_hub_genes_robust_mrna_strict_medium_confidence.csv")
ml_qc = read_csv("results/models/ml_outputs_discovery_only_mrna/ml_triple_qc.csv").iloc[0]
ml_ext = read_csv("results/models/ml_outputs_discovery_only_mrna/ml_model_performance_external_validation_summary.csv")
ml_panel = read_csv("results/models/ml_outputs_discovery_only_mrna/ml_final_biomarker_panel.csv")
pub_comp = read_csv("results/published_signature_validation/ml_vs_published_signature_comparison.csv")
sc_qc = read_csv("results/single_cell_validation/single_cell_validation_triple_qc.csv")
sc_man_qc = read_csv("results/single_cell_validation/single_cell_manuscript_summary_qc.csv").iloc[0]
sc_top = read_csv("results/single_cell_validation/single_cell_clean_top_celltype_gene_changes.csv")
graded_axes = read_csv("results/submission_enhancements/mirna_mrna_axes_evidence_graded.csv")
priority_table = read_csv("results/submission_enhancements/final_target_priority_integrated.csv")
priority_qc = read_csv("results/submission_enhancements/final_target_priority_qc.csv").iloc[0]
ml_interp = read_csv("results/submission_enhancements/ml_external_permutation_importance.csv")
ml_interp_qc = read_csv("results/submission_enhancements/ml_interpretability_qc.csv").iloc[0]
bio_summary = read_csv("results/submission_enhancements/biological_model_summary.csv")
mech_qc = read_csv("results/mechanistic_extension/mechanistic_extension_qc.csv").iloc[0]
module_overlap = read_csv("results/mechanistic_extension/candidate_module_overlap.csv")
lr_top = read_csv("results/mechanistic_extension/top_curated_ligand_receptor_interactions.csv")
vk_proxy = read_csv("results/mechanistic_extension/perturbation_priority_proxy_summary.csv")
tnik_evidence = read_csv("results/mechanistic_extension/tnik_evidence_summary.csv")


bulk = expr_qc[expr_qc["data_type"].str.contains("bulk", case=False, na=False)]
mirna = expr_qc[expr_qc["data_type"].str.contains("miRNA", case=False, na=False)]
bulk_samples = int(bulk["include_yes_samples"].sum())
bulk_ipf = int(bulk["ipf_samples"].sum())
bulk_ctrl = int(bulk["control_samples"].sum())
mirna_samples = int(mirna["include_yes_samples"].sum())
mirna_ipf = int(mirna["ipf_samples"].sum())
mirna_ctrl = int(mirna["control_samples"].sum())
total_bulk_mirna_samples = bulk_samples + mirna_samples
total_bulk_mirna_ipf = bulk_ipf + mirna_ipf
total_bulk_mirna_ctrl = bulk_ctrl + mirna_ctrl

mrna_de = robust_qc[robust_qc["data_type"].eq("bulk mRNA")].iloc[0]
mirna_de = robust_qc[robust_qc["data_type"].eq("miRNA")].iloc[0]
ml_best = ml_ext.sort_values("mean_external_roc_auc", ascending=False).iloc[0]
pub_best = pub_comp[pub_comp["comparator_type"].eq("published_signature")].sort_values("mean_external_roc_auc", ascending=False).iloc[0]

top_mrna = robust_mrna.head(10)["gene_symbol"].dropna().tolist()
top_mirna = robust_mirna.head(10)["mirna_name"].dropna().tolist()
top_panel = ml_panel.head(25)["feature"].tolist()
top_hubs = ppi_hubs.head(12)["gene_symbol"].tolist()
top_terms = enrich_strict.head(6)[["database", "Description", "p.adjust", "Count"]]
top_axes = axes.head(8)[["axis", "match_type", "axis_score"]]
exact_axis_count = int((graded_axes["match_type"] == "exact").sum())
arm_axis_count = int((graded_axes["match_type"] == "arm_agnostic").sum())
main_text_axis_count = int((graded_axes["recommended_manuscript_role"] == "main_text_prioritized_axis").sum())
tier1_genes = priority_table[priority_table["priority_tier"].eq("Tier 1")]["gene_symbol"].head(10).tolist()
top_priority_genes = priority_table.head(10)["gene_symbol"].tolist()
top_perm_features = ml_interp.head(6)["feature"].tolist()

sc_pass = sc_qc[sc_qc["triple_qc_pass"].astype(str).str.lower().eq("true")]
sc_cells = int(sc_pass["metadata_cells_ipf_control"].sum())
sc_target_found = sc_pass["target_genes_found"].astype(int).tolist()

dataset_table = expr_qc[
    [
        "series_id",
        "data_type",
        "dataset_role",
        "features",
        "include_yes_samples",
        "ipf_samples",
        "control_samples",
        "value_scale_guess",
        "triple_qc_pass",
    ]
].copy()
dataset_table["IPF/control"] = dataset_table["ipf_samples"].astype(str) + "/" + dataset_table["control_samples"].astype(str)
dataset_table = dataset_table[
    [
        "series_id",
        "data_type",
        "dataset_role",
        "features",
        "include_yes_samples",
        "IPF/control",
        "value_scale_guess",
        "triple_qc_pass",
    ]
].rename(columns={"include_yes_samples": "samples"})

de_table = de_qc[
    [
        "series_id",
        "data_type",
        "dataset_role",
        "method",
        "tested_features",
        "samples_total",
        "ipf_samples",
        "control_samples",
        "significant_fdr_0_05_logfc_1",
        "top_feature",
        "top_logFC",
        "top_adj_p",
    ]
].copy()
de_table["IPF/control"] = de_table["ipf_samples"].astype(str) + "/" + de_table["control_samples"].astype(str)
de_table["top_adj_p"] = de_table["top_adj_p"].map(fmt_p)
de_table["top_logFC"] = de_table["top_logFC"].map(lambda x: fmt_float(x, 2))
de_table = de_table[
    [
        "series_id",
        "data_type",
        "dataset_role",
        "method",
        "samples_total",
        "IPF/control",
        "significant_fdr_0_05_logfc_1",
        "top_feature",
        "top_logFC",
        "top_adj_p",
    ]
].rename(columns={"significant_fdr_0_05_logfc_1": "sig_FDR0.05_absLogFC1"})

panel_table = ml_panel.head(25)[
    [
        "feature",
        "overall_selection_frequency",
        "discovery_logFC",
        "same_direction_fdr_sig_count",
        "robust_score",
        "axis_count",
        "target_priority_score",
        "oligonucleotide_strategy",
    ]
].copy()
panel_table["overall_selection_frequency"] = panel_table["overall_selection_frequency"].map(lambda x: fmt_float(x, 2))
panel_table["discovery_logFC"] = panel_table["discovery_logFC"].map(lambda x: fmt_float(x, 2))
panel_table["target_priority_score"] = panel_table["target_priority_score"].map(lambda x: fmt_float(x, 2))
panel_table = panel_table[
    [
        "feature",
        "overall_selection_frequency",
        "discovery_logFC",
        "same_direction_fdr_sig_count",
        "axis_count",
        "target_priority_score",
        "oligonucleotide_strategy",
    ]
].rename(columns={"overall_selection_frequency": "selection_frequency"})

sc_table = sc_top.head(12)[
    [
        "series_id",
        "broad_celltype",
        "fine_celltype",
        "gene_symbol",
        "ipf_minus_control_log1p_mean_norm",
        "ipf_minus_control_detection_fraction",
        "sources",
    ]
].copy()
sc_table["ipf_minus_control_log1p_mean_norm"] = sc_table["ipf_minus_control_log1p_mean_norm"].map(lambda x: fmt_float(x, 2))
sc_table["ipf_minus_control_detection_fraction"] = sc_table["ipf_minus_control_detection_fraction"].map(lambda x: fmt_float(x, 2))

priority_display_table = priority_table.head(12)[
    [
        "gene_symbol",
        "final_priority_score",
        "priority_tier",
        "direction",
        "in_ml_panel",
        "suggested_oligonucleotide_strategy",
        "top_single_cell_context",
    ]
].copy()
priority_display_table["final_priority_score"] = priority_display_table["final_priority_score"].map(lambda x: fmt_float(x, 2))

bio_display_table = bio_summary[
    [
        "module",
        "supporting_genes",
        "main_cellular_context",
        "oligonucleotide_angle",
    ]
].copy()

module_display_table = module_overlap.sort_values(["robust_mrna_count", "ml_panel_count"], ascending=False).head(6)[
    [
        "module",
        "module_size",
        "robust_mrna_count",
        "ml_panel_count",
        "priority_top80_count",
        "contains_tnik",
        "ipf_module_r",
        "ipf_module_fdr",
    ]
].copy()
module_display_table["ipf_module_r"] = module_display_table["ipf_module_r"].map(lambda x: fmt_float(x, 2))
module_display_table["ipf_module_fdr"] = module_display_table["ipf_module_fdr"].map(fmt_p)

lr_display_table = lr_top.head(12)[
    [
        "series_id",
        "pair",
        "sender_celltype",
        "receiver_celltype",
        "pathway",
        "ipf_minus_control_interaction_score",
    ]
].copy()
lr_display_table["ipf_minus_control_interaction_score"] = lr_display_table["ipf_minus_control_interaction_score"].map(lambda x: fmt_float(x, 2))

vk_display_table = vk_proxy[
    [
        "target_gene",
        "target_module",
        "gse32537_logFC",
        "gse32537_adj_p",
        "coexpression_neighbors_abs_r_ge_0_35",
        "perturbation_priority_proxy_score",
        "oligonucleotide_interpretation",
        "claim_strength",
    ]
].copy()
vk_display_table["gse32537_logFC"] = vk_display_table["gse32537_logFC"].map(lambda x: fmt_float(x, 2))
vk_display_table["gse32537_adj_p"] = vk_display_table["gse32537_adj_p"].map(fmt_p)
vk_display_table["perturbation_priority_proxy_score"] = vk_display_table["perturbation_priority_proxy_score"].map(lambda x: fmt_float(x, 1))

tnik_display_table = tnik_evidence[
    [
        "evidence_layer",
        "series_id",
        "metric",
        "value",
        "supports_tnik_as_primary_discovery_target",
        "interpretation",
    ]
].copy()


def md_table(df: pd.DataFrame) -> str:
    rows = []
    clean = df.copy()
    for col in clean.columns:
        clean[col] = clean[col].map(lambda x: fmt_float(x) if isinstance(x, float) else str(x))
    rows.append("| " + " | ".join(clean.columns) + " |")
    rows.append("| " + " | ".join(["---"] * len(clean.columns)) + " |")
    for _, row in clean.iterrows():
        rows.append("| " + " | ".join(str(row[col]).replace("|", "/") for col in clean.columns) + " |")
    return "\n".join(rows)


title = (
    "Multi-cohort transcriptomic integration and machine learning identify "
    "oligonucleotide-tractable candidate targets in idiopathic pulmonary fibrosis"
)

abstract = f"""Background: Idiopathic pulmonary fibrosis (IPF) is a progressive fibrotic lung disease with limited therapeutic options. Oligonucleotide therapeutics, including antisense oligonucleotides, siRNAs, and miRNA-directed approaches, require target candidates that are reproducible across cohorts and biologically interpretable. We integrated public mRNA, miRNA, and single-cell transcriptomic data to prioritize robust IPF-associated genes and miRNA-mRNA regulatory axes suitable for downstream oligonucleotide-focused validation.

Results: Eight bulk mRNA or miRNA GEO datasets passed three-layer quality control, comprising {total_bulk_mirna_samples} profiled samples ({total_bulk_mirna_ipf} IPF and {total_bulk_mirna_ctrl} controls). Differential expression in discovery cohorts identified {int(mrna_de['discovery_significant_features'])} significant mRNA genes and {int(mirna_de['discovery_significant_features'])} significant miRNAs after annotation and gene-level harmonization. Cross-cohort direction-consistency screening retained {int(mrna_de['robust_strict_candidates'])} robust mRNA candidates and {int(mirna_de['robust_strict_candidates'])} robust miRNA candidates. miRTarBase integration yielded {int(axis_qc['negative_direction_axes'])} negatively directed miRNA-mRNA candidate axes; evidence grading retained {main_text_axis_count} exact mature-miRNA axes for main-text prioritization and assigned {arm_axis_count} arm-agnostic axes to exploratory supplementary status. Enrichment analysis of robust mRNAs highlighted cilium movement, axoneme assembly, and extracellular matrix-related programs, while STRING network analysis prioritized hubs including {', '.join(top_hubs[:8])}. A leakage-controlled machine-learning workflow using discovery-only mRNA features selected Elastic Net as the best model, with mean external ROC AUC {fmt_float(ml_best['mean_external_roc_auc'], 3)}, minimum external ROC AUC {fmt_float(ml_best['min_external_roc_auc'], 3)}, and mean external PR AUC {fmt_float(ml_best['mean_external_pr_auc'], 3)} across four validation cohorts. External interpretability analysis identified {', '.join(top_perm_features[:4])} as leading permutation-importance features. The resulting 25-gene panel included {', '.join(top_panel[:10])}. Compared with published IPF signatures evaluated under the same external validation framework, the proposed model showed higher mean external ROC AUC ({fmt_float(ml_best['mean_external_roc_auc'], 3)} versus {fmt_float(pub_best['mean_external_roc_auc'], 3)} for the best published comparator). Integrated target prioritization nominated Tier 1 genes including {', '.join(tier1_genes)}. Mechanistic extension identified two IPF-associated coexpression modules, M02 and M09, enriched for robust and machine-learning candidates, and curated ligand-receptor scoring highlighted MIF-CD74, COL14A1-ITGB1, CXCL12-CXCR4, POSTN-ITGB1, and SPP1-integrin communication changes. Single-cell validation in two IPF datasets covering {sc_cells} metadata-matched cells localized key candidates to disease-relevant compartments, including SPP1 in macrophage populations, COL1A1/COL3A1/POSTN in myofibroblasts, and reduced GPX3 in stromal subsets. TNIK showed validation-cohort differential-expression and Wnt/TNIK bridge evidence but was not a discovery-selected or machine-learning-selected primary target.

Conclusions: This multi-cohort computational study prioritizes reproducible IPF-associated mRNA and miRNA candidates and nominates oligonucleotide-tractable axes for functional validation. The findings support an integrated epithelial, stromal, immune, and ciliary remodeling signature in IPF, while emphasizing the need for experimental perturbation and prospective validation before clinical translation."""


sections = []
sections.append(f"# {title}\n")
sections.append("Authors: Yunyi Zhou¹; Yanli Zhang¹*\n\nAffiliation: ¹State Key Laboratory of Common Mechanism Research for Major Diseases, Department of Biochemistry and Molecular Biology, Institute of Basic Medical Sciences, Chinese Academy of Medical Sciences and Peking Union Medical College, Beijing, China.\n\nCorresponding author: Yanli Zhang, zhangyanli@ibms.pumc.edu.cn\n")
sections.append("## Abstract\n" + abstract + "\n")
sections.append("## Keywords\nIdiopathic pulmonary fibrosis; oligonucleotide therapeutics; machine learning; miRNA; transcriptomics; single-cell RNA sequencing; biomarker discovery\n")

sections.append(
    """## Background
Idiopathic pulmonary fibrosis (IPF) is characterized by progressive distortion of lung architecture, aberrant epithelial repair, fibroblast activation, extracellular matrix accumulation, and immune remodeling. Although antifibrotic therapies can slow functional decline, they do not reverse established disease, and there remains a need for mechanistically grounded molecular targets. Oligonucleotide therapeutics provide a direct route to modulate transcripts or miRNA activity, but the translational value of a candidate depends on reproducibility across cohorts, compatibility with target directionality, and evidence that the candidate is active in relevant lung cell populations.

Public transcriptomic resources make it possible to evaluate these criteria systematically. However, many IPF biomarker studies rely on single discovery cohorts, limited validation, or feature selection strategies that risk information leakage when validation data are used before model testing. In addition, mRNA and miRNA evidence is often analyzed separately, making it difficult to nominate coordinated regulatory axes for oligonucleotide intervention.

Here, we constructed a reproducible computational framework that combines GEO mRNA and miRNA datasets, external validation, miRTarBase target evidence, pathway enrichment, STRING protein interaction analysis, machine learning, published-signature comparison, and single-cell localization. The primary objective was to identify robust IPF-associated molecular candidates and prioritize those with plausible oligonucleotide-focused validation routes.
"""
)

sections.append(
    f"""## Results
### Study design and quality-controlled datasets
The workflow integrated bulk mRNA, miRNA, and single-cell transcriptomic data from public GEO datasets (Figure 1). Eight bulk or miRNA datasets passed three independent quality-control checks: annotation completeness, expression-sample cross-matching, and matrix integrity. These datasets contained {bulk_samples} bulk mRNA samples ({bulk_ipf} IPF and {bulk_ctrl} controls) and {mirna_samples} miRNA samples ({mirna_ipf} IPF and {mirna_ctrl} controls). Two single-cell datasets were available for expression-level validation, while GSE122960 was excluded because a cell-level expression matrix was not available locally.

Table 1 summarizes the quality-controlled expression datasets.
"""
)
sections.append(md_table(dataset_table) + "\n")

sections.append(
    f"""### Differential expression and cross-cohort robust candidate selection
Differential expression was performed using limma for log-scale or normalized matrices and edgeR-limma voom for count-like matrices. In the mRNA discovery dataset GSE32537, {int(de_qc.loc[de_qc['series_id'].eq('GSE32537'), 'significant_fdr_0_05_logfc_1'].iloc[0])} features passed FDR < 0.05 and absolute log fold change >= 1; after feature annotation and gene-level harmonization, {int(mrna_de['discovery_significant_features'])} mRNA genes were used for discovery-validation screening. In the miRNA discovery dataset GSE32538, {int(mirna_de['discovery_significant_features'])} significant miRNAs were detected.

Direction-consistency screening across independent validation datasets retained {int(mrna_de['robust_strict_candidates'])} strict robust mRNA candidates and {int(mirna_de['robust_strict_candidates'])} strict robust miRNA candidates. The highest-ranked robust mRNAs included {', '.join(top_mrna)}, while the robust miRNA set included {', '.join(top_mirna)}. These candidates were carried forward for target-axis construction, enrichment, network analysis, and machine learning.
"""
)
sections.append("Table 2 summarizes differential-expression outputs by dataset.\n")
sections.append(md_table(de_table) + "\n")

sections.append(
    f"""### miRNA-mRNA regulatory axes nominate oligonucleotide-relevant candidate relationships
We integrated the robust miRNA and mRNA candidate sets with the human miRTarBase 2025 v10 interaction table. Of {int(axis_qc['robust_mirna_candidates'])} robust miRNAs and {int(axis_qc['robust_mrna_candidates'])} robust mRNAs, {int(axis_qc['matched_mti_rows_for_candidate_mirnas'])} miRTarBase records matched candidate miRNAs. Applying an opposite-direction rule between miRNA and mRNA discovery log fold changes identified {int(axis_qc['negative_direction_axes'])} negative candidate axes involving {int(axis_qc['unique_candidate_mirnas_in_axes'])} miRNAs and {int(axis_qc['unique_target_genes_in_axes'])} target genes. Evidence grading separated {exact_axis_count} exact mature-miRNA matches from {arm_axis_count} arm-agnostic matches. The exact hsa-miR-375 -> CLDN1, hsa-miR-375 -> MNS1, and hsa-miR-375 -> RPGRIP1L axes were retained as main-text prioritized hypotheses. Arm-agnostic matches, including hsa-miR-92a with TP63 and hsa-miR-30a with CDH2, were retained only as exploratory supplementary hypotheses because they require mature-arm-specific validation before therapeutic interpretation.
"""
)
sections.append("Top miRNA-mRNA axes are shown below.\n")
sections.append(md_table(top_axes) + "\n")

sections.append(
    f"""### Pathway enrichment and protein interaction analysis highlight ciliary and matrix remodeling modules
GO, KEGG, and Reactome enrichment analyses used the gene-level feature universe tested in the GSE32537 discovery analysis as background. For the {int(enrich_qc.loc[enrich_qc['gene_set'].eq('robust_mrna_strict'), 'input_symbols'].iloc[0])} strict robust mRNAs, {int(enrich_qc.loc[enrich_qc['gene_set'].eq('robust_mrna_strict'), 'mapped_symbols'].iloc[0])} symbols mapped successfully and {int(enrich_qc.loc[enrich_qc['gene_set'].eq('robust_mrna_strict'), 'total_significant_terms'].iloc[0])} significant enrichment terms were detected. The most significant terms included cilium movement, cilium movement involved in cell motility, microtubule-based movement, and axoneme assembly. Terms containing sperm or flagellar labels were interpreted as shared axonemal/ciliary structural programs rather than reproductive biology.

STRING analysis mapped {int(ppi_qc.loc[(ppi_qc['gene_set'].eq('robust_mrna_strict')) & (ppi_qc['score_label'].eq('medium_confidence')), 'mapped_genes'].iloc[0])} of {int(ppi_qc.loc[(ppi_qc['gene_set'].eq('robust_mrna_strict')) & (ppi_qc['score_label'].eq('medium_confidence')), 'input_genes'].iloc[0])} robust mRNAs at medium confidence, yielding {int(ppi_qc.loc[(ppi_qc['gene_set'].eq('robust_mrna_strict')) & (ppi_qc['score_label'].eq('medium_confidence')), 'edges'].iloc[0])} edges and a largest connected component of {int(ppi_qc.loc[(ppi_qc['gene_set'].eq('robust_mrna_strict')) & (ppi_qc['score_label'].eq('medium_confidence')), 'largest_component_size'].iloc[0])} nodes. The top-ranked hubs were {', '.join(top_hubs)}.
"""
)
sections.append("Representative enriched terms are listed below.\n")
sections.append(md_table(top_terms.assign(**{"p.adjust": top_terms["p.adjust"].map(fmt_p)})) + "\n")

sections.append(
    f"""### Leakage-controlled machine learning identifies a 25-gene IPF panel
To avoid validation leakage, machine learning used only mRNA features selected from the GSE32537 discovery analysis. After harmonizing gene-level matrices across the discovery and external validation cohorts, {int(ml_qc['common_features'])} common discovery-only mRNA features were available. Seven model families were evaluated with imputation, scaling, and SelectKBest feature selection inside scikit-learn pipelines; feature selection and tuning were nested within resampling folds. Elastic Net achieved the highest mean external ROC AUC. Across GSE110147, GSE150910, GSE53845, and GSE92592, the best model had mean external ROC AUC {fmt_float(ml_best['mean_external_roc_auc'], 3)}, minimum external ROC AUC {fmt_float(ml_best['min_external_roc_auc'], 3)}, mean external PR AUC {fmt_float(ml_best['mean_external_pr_auc'], 3)}, and mean balanced accuracy {fmt_float(ml_best['mean_external_balanced_accuracy'], 3)}. A label-permutation control produced a mean AUC of {fmt_float(ml_qc['permutation_control_mean_auc'], 3)}, supporting non-random predictive signal.

The final stable panel contained {int(ml_qc['final_panel_size'])} genes: {', '.join(top_panel)}. Many top-ranked panel genes were also robustly direction-consistent across validation datasets and were assigned oligonucleotide strategy labels such as siRNA or antisense knockdown candidates for upregulated mRNAs. Pooled external interpretability analysis gave a ROC AUC of {fmt_float(ml_interp_qc['pooled_external_auc'], 3)} across {int(ml_interp_qc['external_samples'])} validation samples, with GPX3 as the largest permutation-importance feature followed by {', '.join(top_perm_features[1:6])}. Calibration and decision-curve outputs were generated to support model-review discussions rather than to claim immediate clinical deployability.
"""
)
sections.append("Table 3 lists the 25-gene machine-learning panel.\n")
sections.append(md_table(panel_table) + "\n")

sections.append(
    f"""### Comparison with published IPF signatures
Because most published IPF machine-learning studies do not provide deployable model objects or full coefficients, we evaluated reported gene signatures as fixed external comparator feature sets. Each comparator was trained only in GSE32537 and externally validated under the same framework used for the proposed model. The best published comparator, an explainable machine-learning reported-gene set, achieved mean external ROC AUC {fmt_float(pub_best['mean_external_roc_auc'], 3)}. The proposed discovery-only Elastic Net model achieved higher mean external ROC AUC ({fmt_float(ml_best['mean_external_roc_auc'], 3)}) and a higher minimum external ROC AUC ({fmt_float(ml_best['min_external_roc_auc'], 3)}), suggesting improved cross-cohort stability.
"""
)
sections.append(md_table(pub_comp[["comparator_type", "signature_name", "mean_external_roc_auc", "min_external_roc_auc", "mean_external_pr_auc", "mean_external_balanced_accuracy", "feature_count"]]) + "\n")

sections.append(
    f"""### Single-cell validation localizes candidates to disease-relevant compartments
Single-cell validation was performed in GSE135893 and GSE136831, which together provided {sc_cells} IPF/control metadata-matched cells. The pipeline streamed sparse matrices directly, validated metadata alignment, and filtered non-informative cell labels including Multiplet, Outlier, MT-tRNAs, and CellCycle categories. Of 59 requested target genes, {sc_target_found[0]} were found in GSE135893 and {sc_target_found[1]} were found in GSE136831.

The most prominent cell-type signals localized SPP1 to macrophage or myeloid compartments, COL1A1, COL1A2, COL3A1, and POSTN to stromal myofibroblast compartments, and GPX3 downregulation to fibroblast or myofibroblast populations. These patterns support a disease model in which epithelial/ciliary changes, stromal matrix remodeling, and macrophage activation converge in IPF and provide cellular context for oligonucleotide target prioritization.
"""
)
sections.append(md_table(sc_table) + "\n")

sections.append(
    f"""### Immune-stromal module analysis and single-cell communication extension
To increase mechanistic depth beyond differential expression and prediction, we added an immune/stromal marker-score and WGCNA-like coexpression analysis in the GSE32537 discovery cohort. Ten coexpression modules were detected. M02 contained {int(module_overlap.loc[module_overlap['module'].eq('M02'), 'robust_mrna_count'].iloc[0])} robust mRNA candidates and {int(module_overlap.loc[module_overlap['module'].eq('M02'), 'ml_panel_count'].iloc[0])} machine-learning panel genes and correlated with IPF status (r = {fmt_float(module_overlap.loc[module_overlap['module'].eq('M02'), 'ipf_module_r'].iloc[0], 2)}) and epithelial/ciliary marker scores. M09 contained {int(module_overlap.loc[module_overlap['module'].eq('M09'), 'robust_mrna_count'].iloc[0])} robust mRNA candidates, {int(module_overlap.loc[module_overlap['module'].eq('M09'), 'ml_panel_count'].iloc[0])} machine-learning panel genes, and TNIK; it showed the strongest IPF correlation (r = {fmt_float(module_overlap.loc[module_overlap['module'].eq('M09'), 'ipf_module_r'].iloc[0], 2)}). These modules provide a coexpression-level explanation for why the final candidate set is dominated by epithelial/ciliary and matrix-remodeling biology rather than isolated single-gene effects.

We next performed a curated ligand-receptor scoring analysis using the two available single-cell datasets and a predefined set of IPF-relevant ligand-receptor pairs. After excluding non-informative cell labels such as multiplets and outliers, {int(mech_qc['lr_interactions_scored'])} interaction changes were scored. The strongest IPF-increased interactions involved MIF-CD74 across endothelial, epithelial, and immune compartments; COL14A1-ITGB1 within mesenchymal cells; CXCL12-CXCR4 from mesenchymal to immune cells; POSTN-ITGB1; and SPP1-integrin/CD44 signaling. This supports a model in which epithelial/endothelial inflammatory signaling and immune-mesenchymal matrix signaling jointly reinforce fibrotic remodeling.

A coexpression-neighborhood perturbation-priority proxy was then used to prioritize candidates from their discovery-cohort coexpression neighborhoods. This proxy prioritized CD24, COL14A1, POSTN, and PTGFRN as internally supported knockdown-screening candidates. TNIK had validation-cohort differential-expression and Wnt/TNIK bridge evidence, but it was not significant in the GSE32537 discovery analysis and was not selected by the final machine-learning panel. Therefore, TNIK is best positioned as an externally supplied translational hypothesis to test alongside the primary IPF target modules rather than as a target discovered de novo by this study.
"""
)
sections.append("Table 7 summarizes candidate-enriched coexpression modules.\n")
sections.append(md_table(module_display_table) + "\n")
sections.append("Table 8 lists the top curated ligand-receptor score changes.\n")
sections.append(md_table(lr_display_table) + "\n")
sections.append("Table 9 summarizes coexpression-neighborhood perturbation-priority proxy results.\n")
sections.append(md_table(vk_display_table) + "\n")
sections.append("Table 10 summarizes TNIK evidence and claim strength.\n")
sections.append(md_table(tnik_display_table) + "\n")

sections.append(
    f"""### Integrated target prioritization and biological model
To convert the multi-omic results into a practical shortlist for oligonucleotide-focused follow-up, we built an evidence-weighted prioritization table that combined robust differential expression, machine-learning panel membership, PPI hub status, pathway membership, miRNA-axis evidence, single-cell localization, and direction-compatible oligonucleotide strategy labels. This scoring scheme was used for transparent prioritization only; it was not treated as a causal effect estimate. In total, {int(priority_qc['candidate_genes_scored'])} robust genes were scored, of which {int(priority_qc['tier1_genes'])} were assigned Tier 1 priority. The highest-ranked genes were {', '.join(top_priority_genes)}.

The final biological model organized the findings into five interpretable modules: macrophage/myeloid activation, stromal matrix remodeling, stromal antioxidant/metabolic loss, epithelial/ciliary remodeling, and a higher-confidence hsa-miR-375 target-axis module. This organization provides a clearer bridge between biomarker performance and experimentally testable oligonucleotide hypotheses.
"""
)
sections.append("Table 5 lists the top integrated target-prioritization results.\n")
sections.append(md_table(priority_display_table) + "\n")
sections.append("Table 6 summarizes the final biological interpretation modules.\n")
sections.append(md_table(bio_display_table) + "\n")

sections.append(
    """## Discussion
This study integrates mRNA, miRNA, machine-learning, interaction-network, and single-cell evidence to nominate IPF-associated candidates for oligonucleotide-focused follow-up. Several points strengthen the resulting candidate set. First, mRNA and miRNA candidates were not selected from a single cohort alone; instead, discovery signals were required to show directionally consistent validation across independent datasets. Second, the machine-learning workflow used discovery-only feature prefiltering and nested resampling to reduce leakage. Third, published signatures were evaluated as external comparators under the same validation framework, allowing the proposed panel to be benchmarked against prior literature rather than reported in isolation. Finally, single-cell validation localized key candidates to biologically plausible lung compartments.

The mechanistic extension adds a second layer of support. Rather than treating the machine-learning panel as a black-box biomarker list, marker-score and WGCNA-like module analysis showed that candidate genes concentrate in IPF-associated epithelial/ciliary and matrix-remodeling modules. Curated ligand-receptor scoring then connected these modules to altered intercellular signaling, particularly MIF-CD74, COL14A1-ITGB1, CXCL12-CXCR4, POSTN-ITGB1, and SPP1-integrin/CD44 interactions. These results provide a more coherent biological argument: the nominated targets are embedded in a disease-associated communication network between epithelial, endothelial, immune, and mesenchymal compartments.

The robust mRNA set captured several expected IPF-associated programs. Upregulated matrix and fibroblast-related genes, including ASPN, COL14A1, COL1A1, COL3A1, POSTN, and THY1, were supported by PPI and single-cell evidence. The enrichment of cilium movement and axoneme assembly terms is also notable. Although several ontology labels reference sperm motility, the shared genes encode axonemal and microtubule-associated machinery relevant to motile cilia and epithelial biology. This supports a broader interpretation involving epithelial remodeling and mucociliary dysfunction rather than a reproductive process.

The miRNA analysis suggested several downregulated miRNAs with potential replacement or modulation relevance. hsa-miR-375, hsa-miR-30a, hsa-miR-30d, and hsa-miR-92a were directionally consistent across miRNA validation datasets, and miRTarBase integration connected these miRNAs to upregulated mRNA candidates. We explicitly separated exact mature-miRNA evidence from arm-agnostic evidence because these categories do not carry the same evidentiary weight. The hsa-miR-375 -> CLDN1/MNS1/RPGRIP1L axes were exact mature-miRNA matches and therefore suitable for main-text prioritization, whereas hsa-miR-30a and hsa-miR-92a axes remain exploratory until mature-arm expression and target repression are validated in IPF-relevant cells. These relationships should be considered prioritization hypotheses, because miRNA effects are context-dependent and can involve many targets.

From an oligonucleotide-therapeutic perspective, the final panel contains two broad candidate classes. Upregulated mRNAs with stable validation and cell-type localization may be considered for knockdown strategies such as siRNA or antisense oligonucleotides. Downregulated miRNAs with inverse mRNA target relationships may motivate miRNA mimic or pathway-restoration strategies. The integrated priority table helps distinguish target-like candidates such as COL14A1, ASPN, PTGFRN, CD24, and CDH3 from pathway-state markers such as GPX3, where restoration or pathway protection may be more biologically plausible than direct knockdown. However, delivery to fibrotic lung compartments, target-cell uptake, off-target effects, and disease-stage specificity remain major translational barriers.

TNIK requires cautious interpretation. Public-data evidence placed TNIK in an IPF-associated coexpression module and showed differential expression in several validation cohorts, but it was not significant in the primary discovery cohort and was not selected into the final machine-learning panel. Therefore, a TNIK-targeting small nucleic acid can be connected to this manuscript only as an externally motivated translational bridge, preferably through Wnt/TNIK pathway readouts and co-culture rescue experiments. If the TNIK reagent is a knockdown oligonucleotide, the public IPF transcriptomic direction does not by itself support a simple disease-reversal claim; its value should be tested experimentally in the specific profibrotic cell state where TNIK activity is hypothesized to be pathogenic.

The wet-lab validation path should begin with human lung fibroblasts, alveolar epithelial cells, and macrophage or monocyte-derived macrophage co-cultures. First, confirm baseline expression and perturbation efficiency for COL14A1, POSTN, CD24, PTGFRN, SPP1, and TNIK by qPCR and western blot or targeted proteomics. Second, induce profibrotic stress using TGF-beta1, epithelial injury stimuli, or macrophage-conditioned medium, then test candidate siRNA/ASO or miRNA mimic effects on collagen deposition, alpha-SMA, fibronectin, epithelial injury markers, inflammatory cytokines, and MIF/CD74, SPP1/integrin, CXCL12/CXCR4, and Wnt/TNIK pathway readouts. Third, prioritize combinations where the TNIK reagent modifies Wnt/beta-catenin/TNIK activity and enhances or clarifies the effects of matrix-focused oligonucleotides. These experiments would convert the present computational map into a candidate-validation pipeline.

This study has limitations. It is retrospective and relies on public datasets generated across different platforms, tissues, protocols, and clinical annotation schemes. Although we applied strict sample matching and matrix QC, residual cohort effects cannot be eliminated fully. The machine-learning model was externally validated across public cohorts but not prospectively validated. Published model comparison was based on reported gene signatures rather than original trained model objects, because most prior studies did not provide deployable model files or coefficients. Finally, single-cell validation supports cellular localization but does not prove causal function. Experimental perturbation in relevant human lung cell systems and in vivo models is required before therapeutic interpretation.
"""
)

sections.append(
    """## Conclusions
The integrated analysis identified 280 robust mRNA candidates, 10 robust miRNA candidates, 22 opposite-direction miRNA-mRNA axes, and a 25-gene machine-learning panel with strong external validation. Evidence grading retained three exact hsa-miR-375 target axes as the strongest miRNA-mRNA hypotheses and kept arm-agnostic axes as exploratory signals. Integrated candidate prioritization highlighted Tier 1 candidates including COL14A1, GPX3, ASPN, COL1A1, NECAB1, SPP1, PTGFRN, CD24, POSTN, and CDH3. The evidence converges on ciliary/epithelial remodeling, extracellular matrix activation, macrophage-associated SPP1 signaling, and stromal GPX3 loss as important IPF-associated patterns. These results provide a prioritized, QC-traceable candidate map for downstream oligonucleotide-focused validation.
"""
)

sections.append(
    """## Methods
### Data collection and annotation
Public GEO datasets relevant to human IPF mRNA, miRNA, and single-cell transcriptomics were downloaded and organized locally. Non-human data were excluded. For each bulk or miRNA dataset, sample annotation files were curated with consistent identifiers, disease labels, data type, dataset role, and inclusion status. The expression matrix column identifier was treated as the primary sample_id, while GEO accession identifiers were retained separately when available.

### Expression extraction and three-layer QC
Expression matrices were extracted from GEO series matrix files or supplementary count files. Each dataset underwent three independent QC checks: annotation completeness, sample cross-match between annotation and expression columns, and expression-matrix integrity including numeric coercion, missingness, duplicate feature checks, and scale assessment. Only datasets passing all three checks were used in downstream bulk or miRNA analyses.

### Differential expression
Differential expression was performed in R. Log-scale or normalized intensity matrices were analyzed with limma, while count-like matrices were analyzed with edgeR and limma-voom. The primary contrast was IPF versus control. Discovery significance used FDR < 0.05 and absolute log fold change >= 1. Probe IDs, transcript IDs, and miRNA probe IDs were harmonized to gene symbols or miRNA names before cross-cohort comparison.

### Robust candidate selection
GSE32537 served as the mRNA discovery cohort and GSE32538 as the miRNA discovery cohort. mRNA validation used GSE110147, GSE150910, GSE53845, and GSE92592. miRNA validation used GSE21394 and GSE27430. Strict robust mRNAs required at least two same-direction validation datasets with FDR < 0.05 and absolute log fold change >= 1. Strict robust miRNAs required at least one same-direction validation dataset meeting the miRNA validation rule.

### miRNA-mRNA target-axis integration
Robust miRNAs and mRNAs were integrated with the human miRTarBase 2025 v10 interaction table. Candidate axes were retained when the miRNA and mRNA discovery log fold changes were in opposite directions and the target gene was present in the robust mRNA set. Exact mature-miRNA matches and arm-agnostic matches were annotated separately.

### Evidence grading and integrated target prioritization
miRNA-mRNA axes were graded by match specificity and experimental-support category. Exact mature-miRNA matches were eligible for main-text prioritization, while arm-agnostic matches were retained as exploratory hypotheses requiring mature-arm-specific validation. Robust mRNA candidates were then scored using a transparent additive framework incorporating robust-expression support, machine-learning panel evidence, PPI hub status, enrichment membership, miRNA-axis evidence, single-cell disease-control signal, and oligonucleotide strategy compatibility. The score was used to rank follow-up candidates, not to estimate causal effect size.

### Marker-score modules, single-cell communication, and perturbation proxy
Bulk immune, stromal, epithelial/ciliary, endothelial, and Wnt/TNIK marker scores were computed as within-dataset mean z scores of curated marker genes. In the GSE32537 discovery cohort, highly variable genes plus forced inclusion of robust candidates, machine-learning panel genes, ligand-receptor genes, and TNIK were clustered by correlation distance to generate WGCNA-like coexpression modules. Module eigengenes were calculated by singular-value decomposition and correlated with disease status and marker scores. Single-cell communication potential was approximated by curated ligand-receptor scoring: mean normalized ligand expression in sender cell types was multiplied by mean normalized receptor expression in receiver cell types, and IPF-control score differences were calculated after excluding multiplet, outlier, cell-cycle, and mitochondrial tRNA labels. A coexpression-neighborhood perturbation-priority proxy ranked candidate perturbations by discovery-cohort coexpression-neighborhood structure and target directionality. These analyses are mechanism-generating extensions and should not be interpreted as outputs from full immune-deconvolution, formal cell-communication, or network-knockout packages.

### Enrichment and PPI analysis
GO, KEGG, and Reactome enrichment analyses were performed with clusterProfiler and ReactomePA using the gene-level background tested in the GSE32537 discovery analysis. STRING protein interaction networks were queried for Homo sapiens at medium and high confidence thresholds. Hubs were ranked by combined degree, weighted degree, betweenness, and closeness metrics.

### Machine learning and published-signature comparison
Gene-level matrices were generated by collapsing multiple probes or transcripts to gene symbols using within-dataset variance. The main machine-learning analysis used only discovery-significant mRNA features from GSE32537 that were common across external validation datasets. Seven model families were evaluated: lasso logistic regression, Elastic Net, linear support vector machine, radial-basis support vector machine, random forest, gradient boosting, and a small multilayer perceptron. Imputation, scaling, feature selection, and model fitting were implemented inside scikit-learn pipelines. Repeated stratified cross-validation was used internally, with external validation performed only after model selection. Label permutation was used as a negative control.

Published IPF signatures were evaluated as fixed gene-set comparators because deployable trained model objects or full coefficients were generally unavailable. Each comparator was trained in GSE32537 and externally validated on the same four mRNA validation cohorts.

Model interpretability and utility checks were performed on pooled external validation samples using permutation importance, calibration curves, and decision-curve analysis. These outputs were used to assess whether model predictions were driven by biologically plausible features and whether predicted probabilities had reviewable calibration and threshold-utility behavior.

### Single-cell validation
Single-cell sparse expression matrices were processed by streaming matrix entries to avoid excessive memory use. Candidate target genes were evaluated by cell type and disease group. Cell labels corresponding to Multiplet, Outlier, MT-tRNAs, or CellCycle categories were excluded from manuscript-level summaries, and cell-type comparisons required at least 20 IPF and 20 control cells.

### Software and reproducibility
Analyses were implemented with R, Python, Bioconductor packages including limma, edgeR, clusterProfiler, ReactomePA, org.Hs.eg.db, and Python packages including pandas, numpy, scikit-learn, scipy, matplotlib, and openpyxl. All generated outputs are stored in the project results directory and are traceable through the numbered scripts in the scripts directory.
"""
)

sections.append(
    """## Declarations
### Ethics approval and consent to participate
Not applicable. This study analyzed publicly available de-identified datasets.

### Consent for publication
Not applicable.

### Availability of data and materials
All raw datasets analyzed in this study are publicly available from the Gene Expression Omnibus under the accession numbers reported in Table 1 and the single-cell validation section. Processed non-sensitive outputs supporting the conclusions are included as Additional files 1-10. The numbered analysis scripts used to generate the results are provided as Additional file 11.

### Competing interests
The authors declare that they have no competing interests.

### Funding
No specific funding was received for this study.

### Authors' contributions
Y.Z. curated datasets, performed computational analyses, generated figures and tables, and drafted the manuscript. Y.L.Z. supervised study design, interpreted results, and revised the manuscript. Both authors read and approved the final manuscript.

### Acknowledgements
The authors thank the investigators who generated and shared the public GEO datasets used in this study.

### Use of AI-assisted tools
AI-assisted drafting and coding support were used to organize analysis scripts and prepare an initial manuscript draft. The authors are responsible for verifying all analyses, interpretations, citations, and the final submitted text.
"""
)

sections.append(
    """## Figure and table plan
Figure 1. Overall workflow from GEO data collection, annotation QC, differential expression, robust candidate screening, miRNA-mRNA integration, enrichment/PPI analysis, machine learning, published-signature comparison, and single-cell validation.

Figure 2. Differential expression and robust candidate selection. Suggested panels: discovery volcano plots for GSE32537 and GSE32538; cross-cohort direction matrix for top mRNAs and miRNAs.

Figure 3. miRNA-mRNA axes and enrichment. Suggested panels: miRNA-target network highlighting hsa-miR-375, hsa-miR-30a, and hsa-miR-92a; GO/Reactome dotplot for robust mRNAs.

Figure 4. STRING PPI network and hub prioritization. Suggested panels: robust mRNA PPI network; top hub barplot.

Figure 5. Machine-learning validation. Suggested panels: external ROC curves, external PR curves, feature stability, and comparison with published IPF signatures.

Figure 6. Single-cell localization of prioritized targets. Suggested panels: disease-control delta heatmaps for GSE135893 and GSE136831; selected dotplots for SPP1, COL1A1, POSTN, GPX3, and TPPP3.

Additional Figure S1. Integrated target-prioritization score for the top robust candidates.

Additional Figure S2. Machine-learning permutation importance, calibration, and decision-curve checks in pooled external validation samples.

Additional Figure S3. Biological interpretation model linking robust transcriptomic signals, cell context, and oligonucleotide development hypotheses.

Additional Figure S4. WGCNA-like module-trait heatmap showing IPF-associated epithelial/ciliary and matrix-remodeling modules.

Additional Figure S5. Curated ligand-receptor score changes in IPF single-cell datasets.

Additional Figure S6. Coexpression-neighborhood perturbation-priority ranking for oligonucleotide follow-up candidates including TNIK.

Table 1. Quality-controlled datasets.
Table 2. Differential-expression summary.
Table 3. Final 25-gene machine-learning panel and oligonucleotide strategy labels.
Table 4. Published-signature comparator performance.
Table 5. Integrated target-prioritization shortlist.
Table 6. Biological interpretation modules.
Table 7. Candidate-enriched coexpression modules.
Table 8. Curated ligand-receptor score changes.
Table 9. Coexpression-neighborhood perturbation-priority proxy summary.
Table 10. TNIK evidence and claim-strength table.
Supplementary Table 1. Full merged annotation and matrix QC.
Supplementary Table 2. Full differential expression results.
Supplementary Table 3. Robust mRNA and miRNA candidates.
Supplementary Table 4. miRNA-mRNA axes.
Supplementary Table 5. Enrichment and PPI outputs.
Supplementary Table 6. Machine-learning predictions, performance, and feature stability.
Supplementary Table 7. Single-cell target localization.
Supplementary Table 8. Evidence-graded miRNA-mRNA axes and integrated target-prioritization outputs.
Supplementary Table 9. Mechanistic extension, curated ligand-receptor scoring, perturbation proxy, and TNIK bridge outputs.
"""
)

sections.append(
    """## References to finalize
1. BMC Genomics research article preparation guidance. https://bmcgenomics.biomedcentral.com/submission-guidelines/preparing-your-manuscript/research-article
2. GEO: Gene Expression Omnibus. https://www.ncbi.nlm.nih.gov/geo/
3. limma and edgeR Bioconductor documentation and associated method papers.
4. clusterProfiler, ReactomePA, org.Hs.eg.db, and KEGGREST documentation and associated method papers.
5. STRING database documentation and associated method paper.
6. miRTarBase 2025 v10 human MTI resource.
7. Published IPF comparator signatures: metabolism-related hub genes, ANN six-gene IPF signature, Frontiers hub-gene analysis, explainable ML pulmonary fibrosis biomarkers, MMP1/MMP7 biomarker study, and CXCL14 SHAP report.
8. Finalize primary references for each GEO dataset before submission.
"""
)


manuscript_md = "\n".join(sections)
MD_PATH.write_text(manuscript_md, encoding="utf-8")


def set_cell_text(cell, text: object, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(str(text))
    run.font.size = Pt(size)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, df: pd.DataFrame, title_text: str, font_size: int = 8) -> None:
    doc.add_paragraph(title_text, style="Caption")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, col in enumerate(df.columns):
        set_cell_text(hdr[idx], col, size=font_size)
        shade_cell(hdr[idx], "F4F6F9")
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for idx, col in enumerate(df.columns):
            val = row[col]
            if isinstance(val, float):
                val = fmt_float(val)
            set_cell_text(cells[idx], val, size=font_size)
    doc.add_paragraph()


def add_md_section_to_doc(doc: Document, text: str) -> None:
    for raw in text.splitlines():
        line = raw.strip()
        if not line:
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:], level=0)
        elif line.startswith("## "):
            doc.add_heading(line[3:], level=1)
        elif line.startswith("### "):
            doc.add_heading(line[4:], level=2)
        elif line.startswith("Figure ") or line.startswith("Table ") or line.startswith("Supplementary Table "):
            doc.add_paragraph(line, style="List Bullet")
        elif line.startswith("|"):
            continue
        else:
            doc.add_paragraph(line)


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
styles["Normal"].font.name = "Calibri"
styles["Normal"].font.size = Pt(11)
styles["Normal"].paragraph_format.space_after = Pt(8)
styles["Normal"].paragraph_format.line_spacing = 1.25
for style_name, size, color in [
    ("Title", 16, "0B2545"),
    ("Heading 1", 16, "2E74B5"),
    ("Heading 2", 13, "2E74B5"),
    ("Heading 3", 12, "1F4D78"),
]:
    style = styles[style_name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
styles["Caption"].font.name = "Calibri"
styles["Caption"].font.size = Pt(9)
styles["Caption"].font.italic = True

doc.add_heading(title, level=0)
for p in [
    "Authors: Yunyi Zhou¹; Yanli Zhang¹*",
    "Affiliation: ¹State Key Laboratory of Common Mechanism Research for Major Diseases, Department of Biochemistry and Molecular Biology, Institute of Basic Medical Sciences, Chinese Academy of Medical Sciences and Peking Union Medical College, Beijing, China.",
    "Corresponding author: Yanli Zhang, zhangyanli@ibms.pumc.edu.cn",
]:
    para = doc.add_paragraph(p)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT

doc.add_heading("Abstract", level=1)
for para in abstract.split("\n\n"):
    doc.add_paragraph(para)

doc.add_heading("Keywords", level=1)
doc.add_paragraph("Idiopathic pulmonary fibrosis; oligonucleotide therapeutics; machine learning; miRNA; transcriptomics; single-cell RNA sequencing; biomarker discovery")

doc.add_section(WD_SECTION.NEW_PAGE)

body_without_tables = manuscript_md.split("## Background", 1)[1]
for heading in ["Table 1 summarizes", "Table 2 summarizes", "Top miRNA-mRNA axes", "Representative enriched terms", "Table 3 lists", "Table 4", "### Single-cell validation"]:
    body_without_tables = body_without_tables.replace(heading, "\n" + heading)

parts = body_without_tables.split("Table 1 summarizes the quality-controlled expression datasets.")
add_md_section_to_doc(doc, "## Background\n" + parts[0])
doc.add_paragraph("Table 1 summarizes the quality-controlled expression datasets.")
add_table(doc, dataset_table, "Table 1. Quality-controlled bulk mRNA and miRNA datasets.", 7)

rest = parts[1]
before_de, rest = rest.split("Table 2 summarizes differential-expression outputs by dataset.")
add_md_section_to_doc(doc, before_de)
doc.add_paragraph("Table 2 summarizes differential-expression outputs by dataset.")
add_table(doc, de_table, "Table 2. Differential-expression outputs by dataset.", 7)

before_axes, rest = rest.split("Top miRNA-mRNA axes are shown below.")
add_md_section_to_doc(doc, before_axes)
doc.add_paragraph("Top miRNA-mRNA axes are shown below.")
add_table(doc, top_axes, "Selected high-scoring miRNA-mRNA axes.", 8)

before_terms, rest = rest.split("Representative enriched terms are listed below.")
add_md_section_to_doc(doc, before_terms)
doc.add_paragraph("Representative enriched terms are listed below.")
add_table(doc, top_terms.assign(**{"p.adjust": top_terms["p.adjust"].map(fmt_p)}), "Representative significant enrichment terms.", 8)

before_panel, rest = rest.split("Table 3 lists the 25-gene machine-learning panel.")
add_md_section_to_doc(doc, before_panel)
doc.add_paragraph("Table 3 lists the 25-gene machine-learning panel.")
add_table(doc, panel_table, "Table 3. Final 25-gene machine-learning panel.", 6)

before_pub, rest = rest.split("### Comparison with published IPF signatures")
add_md_section_to_doc(doc, before_pub)
doc.add_heading("Comparison with published IPF signatures", level=2)
before_sc, rest = rest.split("### Single-cell validation localizes candidates to disease-relevant compartments")
add_md_section_to_doc(doc, before_sc)
add_table(doc, pub_comp[["comparator_type", "signature_name", "mean_external_roc_auc", "min_external_roc_auc", "mean_external_pr_auc", "mean_external_balanced_accuracy", "feature_count"]], "Table 4. Proposed model versus published IPF signature comparators.", 7)

doc.add_heading("Single-cell validation localizes candidates to disease-relevant compartments", level=2)
single_cell_text, rest_after_sc_table = rest.split("| series_id |", 1)
add_md_section_to_doc(doc, single_cell_text)
add_table(doc, sc_table, "Selected single-cell disease-control target changes.", 7)
after_sc_table = rest_after_sc_table.split("\n\n", 1)[1]
add_md_section_to_doc(doc, after_sc_table)

doc.save(DOCX_PATH)
print(f"Wrote {MD_PATH}")
print(f"Wrote {DOCX_PATH}")
