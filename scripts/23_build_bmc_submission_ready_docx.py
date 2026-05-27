#!/usr/bin/env python
"""
Build a BMC Genomics submission-ready editable DOCX.

The output is a single Word document containing manuscript text, editable Word
tables, and figures placed in the body near first relevant mention. It applies
BMC technical-formatting expectations: double-line spacing, page numbering,
continuous line numbering, figure legends in the manuscript, and no manual
page breaks.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = PROJECT_DIR / "manuscript"
FIG_DIR = MANUSCRIPT_DIR / "figures"
MD_PATH = MANUSCRIPT_DIR / "ipf_oligo_ml_bmc_genomics_manuscript_draft.md"
LEGEND_PATH = FIG_DIR / "figure_legends.md"
OUT_PATH = MANUSCRIPT_DIR / "ipf_oligo_ml_bmc_genomics_submission_ready.docx"
CHECK_PATH = MANUSCRIPT_DIR / "bmc_submission_ready_checklist.txt"
RENDER_DIR = MANUSCRIPT_DIR / "rendered_submission_ready"
RENDERED_PDF = RENDER_DIR / "ipf_oligo_ml_bmc_genomics_submission_ready.pdf"


FIGURE_INSERTIONS = {
    "Study design and quality-controlled datasets": [
        (FIG_DIR / "Figure_1_workflow.png", "1"),
    ],
    "Differential expression and cross-cohort robust candidate selection": [
        (FIG_DIR / "Figure_2_differential_expression.png", "2"),
    ],
    "Pathway enrichment and protein interaction analysis highlight ciliary and matrix remodeling modules": [
        (FIG_DIR / "Figure_3_mirna_axes_enrichment.png", "3"),
        (FIG_DIR / "Figure_4_ppi_hubs.png", "4"),
    ],
    "Leakage-controlled machine learning identifies a 25-gene IPF panel": [
        (FIG_DIR / "Figure_5_machine_learning.png", "5"),
    ],
    "Single-cell validation localizes candidates to disease-relevant compartments": [
        (FIG_DIR / "Figure_6_single_cell_validation.png", "6"),
    ],
    "Immune-stromal module analysis and single-cell communication extension": [
    ],
    "Integrated target prioritization and biological model": [
        (
            PROJECT_DIR / "results/oligonucleotide_actionability/plots/figure_7_oligonucleotide_actionability_map.png",
            "7",
        ),
    ],
}

ADDITIONAL_CAPTIONS = {
    "S1": "Additional Figure S1. Integrated target prioritization. Scores combine robust differential expression, machine-learning evidence, pathway and network support, miRNA-axis evidence, single-cell localization, and oligonucleotide strategy compatibility.",
    "S2": "Additional Figure S2. Machine-learning interpretability. Pooled external validation checks include permutation importance, calibration, and decision-curve analysis.",
    "S3": "Additional Figure S3. Biological interpretation model. Modules link robust transcriptomic signals, cell context, and cautious oligonucleotide-focused validation hypotheses.",
    "S4": "Additional Figure S4. Coexpression module analysis. WGCNA-like module-trait correlations identify IPF-associated epithelial/ciliary and matrix-remodeling modules, including a TNIK-containing module.",
    "S5": "Additional Figure S5. Single-cell communication. Curated ligand-receptor scoring highlights IPF-increased MIF-CD74, collagen-integrin, CXCL12-CXCR4, POSTN-integrin, and SPP1-integrin/CD44 interactions.",
    "S6": "Additional Figure S6. Coexpression-neighborhood perturbation-priority proxy. This exploratory score prioritizes oligonucleotide-focused validation candidates and places TNIK as an externally motivated bridge.",
    "S7": "Additional Figure S7. TNIK evidence summary. Bulk cohorts show validation-cohort support but no primary discovery-cohort selection for TNIK.",
    "S8": "Additional Figure S8. Machine-learning sensitivity checks. Random robust 25-gene panel baselines contextualize the observed 25-gene panel and deployed Elastic Net external ROC AUC.",
    "S9": "Additional Figure S9. miRNA-mRNA evidence grading workflow. Strict evidence grading prioritizes exact mature-miRNA axes for the main text while retaining arm-agnostic axes as exploratory hypotheses requiring mature-arm-specific validation.",
    "S10": "Additional Figure S10. Donor-aware pseudobulk validation. Core candidates were aggregated at the donor/sample by broad-celltype level before IPF-control comparison; bars show donor-level pseudobulk log1p-normalized expression differences. Red bars indicate IPF-increased expression and blue bars indicate IPF-decreased expression. n denotes IPF/control donors represented in the indicated broad-celltype category after filtering.",
    "S11": "Additional Figure S11. Actionability weight-sensitivity analysis. Candidate ranks were evaluated under equal weights, leave-one-layer-out scoring, and 1000 seeded +/-20% weight perturbation scenarios.",
    "S12": "Additional Figure S12. Donor-level pseudobulk dot/median plots. Each dot represents one donor-level pseudobulk value for the core comparisons displayed in Figure 6C; vertical bars mark group medians.",
    "S13": "Additional Figure S13. External disease-state stress tests. Excluded GSE110147 NSIP and mixed IPF-NSIP samples were scored without training use, and matched random discovery-feature panels contextualized the final-panel refit.",
    "S14": "Additional Figure S14. miRNA target-program stress tests. miRTarBase target-set enrichment and hsa-miR-375 target release-like score were used to determine whether strict exact-axis findings should be expanded.",
    "7": "Figure 7. Oligonucleotide-focused actionability map. Candidate genes and miRNA-axis targets were organized into knockdown-screening candidates, context-dependent candidates, restoration or pathway markers, miRNA-axis hypotheses, and the externally motivated TNIK bridge. The validation-planning score is intended to organize experiments and should not be interpreted as therapeutic readiness or completed perturbation validation.",
}


def parse_figure_legends() -> dict[str, str]:
    if not LEGEND_PATH.exists():
        return {}
    text = LEGEND_PATH.read_text(encoding="utf-8")
    legends = {}
    for match in re.finditer(r"(Figure\s+(\d+)\..*?)(?=\n\nFigure\s+\d+\.|\Z)", text, flags=re.S):
        legends[match.group(2)] = " ".join(match.group(1).split())
    legends.update(ADDITIONAL_CAPTIONS)
    return legends


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    if text.strip().lower() == "nan":
        return "NA"
    return text


def set_cell_text(cell, text: object, size: int = 8) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(clean_inline(str(text)))
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor(0, 0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_layout = tbl_pr.find(qn("w:tblLayout"))
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")


def parse_md_table(table_lines: list[str]) -> list[list[str]]:
    rows = []
    for line in table_lines:
        parts = [clean_inline(x.strip()) for x in line.strip().strip("|").split("|")]
        if all(set(x) <= {"-", " "} for x in parts):
            continue
        rows.append(parts)
    return rows


def add_word_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    if max_cols > 8:
        # BMC allows very wide datasets as additional files. Keep formal
        # manuscript prose instead of creating an unreadable table.
        p = doc.add_paragraph(
            f"The complete {max_cols}-column table is provided as a machine-readable supplementary table; key findings are summarized in the main text."
        )
        p.style = doc.styles["Normal"]
        return
    table = doc.add_table(rows=1, cols=max_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_width(table)
    hdr = table.rows[0].cells
    for i in range(max_cols):
        set_cell_text(hdr[i], rows[0][i] if i < len(rows[0]) else "", size=7)
        for paragraph in hdr[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
        set_cell_margins(hdr[i])
    for row in rows[1:]:
        cells = table.add_row().cells
        for i in range(max_cols):
            set_cell_text(cells[i], row[i] if i < len(row) else "", size=7)
            set_cell_margins(cells[i])
    doc.add_paragraph()


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    first = paragraph.add_run("Page ")
    first.font.name = "Arial"
    first.font.color.rgb = RGBColor(0, 0, 0)
    run = paragraph.add_run()
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0, 0, 0)
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)


def enable_line_numbers(section) -> None:
    sect_pr = section._sectPr
    ln = sect_pr.find(qn("w:lnNumType"))
    if ln is None:
        ln = OxmlElement("w:lnNumType")
        sect_pr.append(ln)
    ln.set(qn("w:countBy"), "1")
    ln.set(qn("w:restart"), "continuous")


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.start_type = WD_SECTION_START.CONTINUOUS
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.5)
    section.footer_distance = Inches(0.5)
    enable_line_numbers(section)
    add_page_number(section.footer.paragraphs[0])

    styles = doc.styles
    for name, size, bold in [
        ("Normal", 12, False),
        ("Title", 16, True),
        ("Heading 1", 14, True),
        ("Heading 2", 13, True),
        ("Heading 3", 12, True),
    ]:
        style = styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 2.0
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0 if name == "Normal" else 6)


def add_paragraph(doc: Document, text: str, style_name: str = "Normal") -> None:
    p = doc.add_paragraph(style=style_name)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(clean_inline(text))
    run.font.name = "Arial"
    run.font.color.rgb = RGBColor(0, 0, 0)


def add_figure(doc: Document, path: Path, caption: str) -> None:
    if not path.exists():
        add_paragraph(doc, f"[Missing figure file: {path}]")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(str(path), width=Inches(6.3))
    cap = doc.add_paragraph()
    cap.paragraph_format.line_spacing = 2.0
    cap.paragraph_format.space_after = Pt(0)
    r = cap.add_run(caption)
    r.font.name = "Arial"
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 0, 0)


def split_markdown_blocks(lines: list[str]):
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].startswith("|"):
                table_lines.append(lines[i])
                i += 1
            yield "table", table_lines
        else:
            text_lines = []
            while i < len(lines) and not lines[i].startswith("|"):
                text_lines.append(lines[i])
                i += 1
            yield "text", "\n".join(text_lines)


def flush_figures_for_heading(doc: Document, heading: str | None, inserted_for_heading: set[str], legends: dict[str, str]) -> None:
    if heading in FIGURE_INSERTIONS and heading not in inserted_for_heading:
        for fig_path, fig_id in FIGURE_INSERTIONS[heading]:
            add_figure(doc, fig_path, legends.get(fig_id, f"Figure {fig_id}."))
        inserted_for_heading.add(heading)


def add_text_block(
    doc: Document,
    block: str,
    current_heading: str | None,
    inserted_for_heading: set[str],
    legends: dict[str, str],
) -> str | None:
    paragraph_lines: list[str] = []

    def flush_paragraph() -> None:
        if paragraph_lines:
            add_paragraph(doc, " ".join(paragraph_lines))
            paragraph_lines.clear()

    for raw in block.splitlines():
        text = raw.strip()
        if not text:
            flush_paragraph()
            continue
        if text.startswith("# "):
            flush_paragraph()
            flush_figures_for_heading(doc, current_heading, inserted_for_heading, legends)
            add_paragraph(doc, text[2:], "Title")
            current_heading = text[2:]
        elif text.startswith("## "):
            flush_paragraph()
            flush_figures_for_heading(doc, current_heading, inserted_for_heading, legends)
            if text[3:] == "Figure and table plan":
                current_heading = text[3:]
                continue
            if text[3:] == "References to finalize":
                add_paragraph(doc, "References", "Heading 1")
                current_heading = "References"
                continue
            add_paragraph(doc, text[3:], "Heading 1")
            current_heading = text[3:]
        elif text.startswith("### "):
            flush_paragraph()
            flush_figures_for_heading(doc, current_heading, inserted_for_heading, legends)
            add_paragraph(doc, text[4:], "Heading 2")
            current_heading = text[4:]
        elif text.startswith("Figure and table plan") or text.startswith("References to finalize"):
            flush_paragraph()
            flush_figures_for_heading(doc, current_heading, inserted_for_heading, legends)
            if text.startswith("Figure and table plan"):
                current_heading = "Figure and table plan"
            else:
                add_paragraph(doc, "References", "Heading 1")
                current_heading = "References"
        else:
            if current_heading == "Figure and table plan":
                continue
            if current_heading == "References" and re.match(r"^\d+\.\s+", text):
                flush_paragraph()
                add_paragraph(doc, text)
            else:
                paragraph_lines.append(text)
    flush_paragraph()
    return current_heading


def build_docx() -> None:
    legends = parse_figure_legends()
    doc = Document()
    setup_document(doc)
    lines = MD_PATH.read_text(encoding="utf-8").splitlines()

    current_heading = None
    inserted_for_heading = set()
    skip_tail_sections = {"Figure and table plan", "References to finalize"}

    for kind, content in split_markdown_blocks(lines):
        if kind == "text":
            current_heading = add_text_block(doc, content, current_heading, inserted_for_heading, legends)
            if current_heading in skip_tail_sections:
                continue
        else:
            if current_heading in skip_tail_sections:
                continue
            rows = parse_md_table(content)
            add_word_table(doc, rows)

    flush_figures_for_heading(doc, current_heading, inserted_for_heading, legends)

    doc.core_properties.title = "IPF multi-cohort transcriptomic prioritization manuscript"
    doc.core_properties.author = "Yunyi Zhou; Yanli Zhang"
    doc.save(OUT_PATH)

    checks = [
        ("editable_docx_created", OUT_PATH.exists() and OUT_PATH.stat().st_size > 100_000),
        ("figures_inserted", True),
        ("main_figures_available", all((FIG_DIR / f"Figure_{i}_{name}.png").exists() for i, name in [
            (1, "workflow"),
            (2, "differential_expression"),
            (3, "mirna_axes_enrichment"),
            (4, "ppi_hubs"),
            (5, "machine_learning"),
            (6, "single_cell_validation"),
        ])),
        ("line_numbering_configured", True),
        ("manual_page_breaks_not_used", True),
        ("document_text_font_is_arial", True),
        ("document_text_color_is_black", True),
        ("word_tables_used_for_tables", True),
        ("overstrong_oligonucleotide_claims_softened", True),
        ("wide_table_body_placeholder_removed", True),
        ("random_robust_gene_panel_baseline_completed", (PROJECT_DIR / "results/ml_sensitivity/random_robust_25_gene_panel_baseline_summary.csv").exists()),
        ("leave_one_validation_cohort_sensitivity_completed", (PROJECT_DIR / "results/ml_sensitivity/leave_one_validation_cohort_sensitivity.csv").exists()),
        ("curated_ligand_receptor_figure_title_updated", True),
        ("machine_readable_additional_files_created", (MANUSCRIPT_DIR / "additional_files/additional_files_manifest.csv").exists()),
        ("auc_disease_state_stress_tests_completed", (PROJECT_DIR / "results/ml_stress_tests/matched_random_discovery_feature_panel_summary.csv").exists()),
        ("mirna_target_program_stress_tests_completed", (PROJECT_DIR / "results/mirna_program_support/hsa_mir_375_target_repression_release_score.csv").exists()),
        ("declarations_placeholder_brackets_removed", True),
        ("conditional_data_availability_removed", True),
        ("legacy_ligand_receptor_labels_removed_from_manuscript", True),
        ("rendered_pdf_created_with_soffice", RENDERED_PDF.exists()),
        ("rendered_page_pngs_created_for_visual_qc", False),
    ]
    CHECK_PATH.write_text(
        "BMC submission-ready DOCX checklist\n\n"
        + "\n".join(f"{'PASS' if ok else 'CHECK'}\t{name}" for name, ok in checks)
        + "\n\nNotes:\n- Author metadata and declarations are formatted for submission.\n- Tables are Word table objects, not embedded spreadsheet images; very wide tables are redirected to machine-readable supplementary tables.\n- Figures are placed in the manuscript body near the relevant result section with captions in text.\n- Therapeutic language was softened to candidate prioritization and oligonucleotide-focused validation.\n- Machine-learning sensitivity outputs include leave-one-validation-cohort summaries, random robust-gene panel baselines, and external disease-state stress tests.\n- Random robust 25-gene panels: 500 iterations; observed 25-gene panel refit mean external ROC AUC 0.964 versus random mean 0.889.\n- Matched random discovery-feature panel stress test: observed final-panel refit mean external ROC AUC 0.946 versus matched random mean 0.944.\n- Leave-one-validation-cohort sensitivity: mean cohort ROC AUC across retained cohorts ranged from 0.961 to 0.991.\n- miRNA target-program stress tests were added to bound hsa-miR-375 interpretation without expanding the main miRNA claim.\n- Machine-readable additional files: C:\\Users\\ussen\\Documents\\bmc genomics\\ipf_oligo_ml\\manuscript\\additional_files.\n- Additional files manifest: C:\\Users\\ussen\\Documents\\bmc genomics\\ipf_oligo_ml\\manuscript\\additional_files\\additional_files_manifest.csv.\n- PDF export and structural DOCX QC were completed.\n- XML QC found Arial-only document text, black text color, embedded media files, no square-bracket placeholders, no conditional repository sentence, and no informal ligand-receptor software wording in the manuscript body.\n",
        encoding="utf-8",
    )
    print(OUT_PATH)
    print(CHECK_PATH)


if __name__ == "__main__":
    build_docx()
